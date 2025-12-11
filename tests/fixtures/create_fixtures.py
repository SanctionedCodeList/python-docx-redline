"""Script to create test fixtures for BytesIO and python-docx compatibility tests."""

from pathlib import Path

from docx import Document

FIXTURES_DIR = Path(__file__).parent


def create_simple_document() -> None:
    """Create a simple test document."""
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a simple test document for testing BytesIO support.")
    doc.add_paragraph("It has multiple paragraphs with some content to search for.")
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.save(FIXTURES_DIR / "simple_document.docx")


if __name__ == "__main__":
    create_simple_document()
    print("Fixtures created!")
