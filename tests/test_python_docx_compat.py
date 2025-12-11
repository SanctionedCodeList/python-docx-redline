"""Tests for python-docx compatibility.

These tests verify the from_python_docx() and to_python_docx() functions
that enable seamless integration between python-docx and python_docx_redline.

Note: These tests are skipped if python-docx is not installed.
"""

from pathlib import Path

import pytest

# Skip all tests if python-docx not installed
docx = pytest.importorskip("docx")
from docx import Document as PythonDocxDocument  # noqa: E402

from python_docx_redline import Document, from_python_docx, to_python_docx  # noqa: E402


class TestFromPythonDocx:
    """Test from_python_docx() conversion function."""

    def test_basic_conversion(self) -> None:
        """Can create python_docx_redline Document from python-docx Document."""
        # Create with python-docx
        py_doc = PythonDocxDocument()
        py_doc.add_paragraph("Hello World")

        # Convert
        doc = from_python_docx(py_doc)

        assert "Hello World" in doc.get_text()

    def test_preserves_paragraphs(self) -> None:
        """Paragraphs are preserved during conversion."""
        py_doc = PythonDocxDocument()
        py_doc.add_paragraph("First paragraph")
        py_doc.add_paragraph("Second paragraph")
        py_doc.add_paragraph("Third paragraph")

        doc = from_python_docx(py_doc)

        assert len(doc.paragraphs) >= 3
        text = doc.get_text()
        assert "First paragraph" in text
        assert "Second paragraph" in text
        assert "Third paragraph" in text

    def test_preserves_headings(self) -> None:
        """Headings are preserved during conversion."""
        py_doc = PythonDocxDocument()
        py_doc.add_heading("Main Title", level=0)
        py_doc.add_heading("Section One", level=1)
        py_doc.add_paragraph("Content under section one.")
        py_doc.add_heading("Section Two", level=1)
        py_doc.add_paragraph("Content under section two.")

        doc = from_python_docx(py_doc)
        text = doc.get_text()

        assert "Main Title" in text
        assert "Section One" in text
        assert "Section Two" in text

    def test_can_make_tracked_changes(self) -> None:
        """Can make tracked edits after converting from python-docx."""
        py_doc = PythonDocxDocument()
        py_doc.add_paragraph("Payment due in 30 days")

        doc = from_python_docx(py_doc)
        doc.replace_tracked("30 days", "45 days")

        assert doc.has_tracked_changes()
        assert "45 days" in doc.get_text()

    def test_custom_author(self) -> None:
        """Can specify custom author during conversion."""
        py_doc = PythonDocxDocument()
        py_doc.add_paragraph("Test content")

        doc = from_python_docx(py_doc, author="Legal Team")
        doc.insert_tracked(" [reviewed]", after="Test content")

        # The author should be set on the document
        assert doc.author == "Legal Team"

    def test_path_is_none(self) -> None:
        """Converted document has no path (in-memory)."""
        py_doc = PythonDocxDocument()
        py_doc.add_paragraph("Test")

        doc = from_python_docx(py_doc)

        assert doc.path is None


class TestFromPythonDocxErrors:
    """Test error handling in from_python_docx()."""

    def test_wrong_type_raises_typeerror(self) -> None:
        """Passing wrong type raises TypeError."""
        with pytest.raises(TypeError, match="Expected python-docx Document"):
            from_python_docx("not a document")

    def test_none_raises_typeerror(self) -> None:
        """Passing None raises TypeError."""
        with pytest.raises(TypeError, match="Expected python-docx Document"):
            from_python_docx(None)

    def test_dict_raises_typeerror(self) -> None:
        """Passing dict raises TypeError."""
        with pytest.raises(TypeError, match="Expected python-docx Document"):
            from_python_docx({"key": "value"})


class TestToPythonDocx:
    """Test to_python_docx() conversion function."""

    def test_basic_conversion(self, tmp_path: Path) -> None:
        """Can convert python_docx_redline Document back to python-docx."""
        # Create a simple document
        py_doc_original = PythonDocxDocument()
        py_doc_original.add_paragraph("Original content")

        # Convert to python_docx_redline
        doc = from_python_docx(py_doc_original)
        doc.replace_tracked("Original", "Modified")

        # Convert back to python-docx (skip validation for in-memory docs)
        py_doc = to_python_docx(doc, validate=False)

        # Should be a python-docx Document (check it has paragraphs attribute)
        assert hasattr(py_doc, "paragraphs")
        # Should have at least one paragraph with content
        assert len(py_doc.paragraphs) >= 1
        # Note: python-docx may not include tracked change text in .text property,
        # so we just verify the document is readable and has paragraphs

    def test_can_continue_editing_with_python_docx(self, tmp_path: Path) -> None:
        """Can add content with python-docx after conversion."""
        py_doc_original = PythonDocxDocument()
        py_doc_original.add_paragraph("Starting text")

        # Make tracked changes
        doc = from_python_docx(py_doc_original)
        doc.insert_tracked(" [tracked]", after="Starting")

        # Convert back and add more content (skip validation for in-memory docs)
        py_doc = to_python_docx(doc, validate=False)
        py_doc.add_paragraph("Added with python-docx")

        # Save and verify
        output_path = tmp_path / "final.docx"
        py_doc.save(str(output_path))

        # Reload and check
        reloaded = PythonDocxDocument(str(output_path))
        full_text = "\n".join(p.text for p in reloaded.paragraphs)
        assert "Added with python-docx" in full_text


class TestFullWorkflow:
    """Test complete workflows combining python-docx and python_docx_redline."""

    def test_create_edit_save_workflow(self, tmp_path: Path) -> None:
        """Full workflow: create with python-docx, edit with python_docx_redline, save."""
        # Step 1: Create document with python-docx
        py_doc = PythonDocxDocument()
        py_doc.add_heading("Contract", 0)
        py_doc.add_paragraph("The payment terms are net 30 days.")
        py_doc.add_paragraph("Effective date: January 1, 2025.")

        # Step 2: Convert and add tracked changes
        doc = from_python_docx(py_doc, author="Legal Team")
        doc.replace_tracked("net 30 days", "net 45 days")
        doc.insert_tracked(" (as amended)", after="Contract")

        # Step 3: Save (skip validation for in-memory docs)
        output_path = tmp_path / "contract_redlined.docx"
        doc.save(output_path, validate=False)

        # Verify
        assert output_path.exists()
        reopened = Document(output_path)
        assert reopened.has_tracked_changes()
        text = reopened.get_text()
        assert "net 45 days" in text
        assert "(as amended)" in text

    def test_in_memory_workflow(self) -> None:
        """Workflow entirely in memory without filesystem."""
        # Create document
        py_doc = PythonDocxDocument()
        py_doc.add_paragraph("Original text here")

        # Add tracked changes
        doc = from_python_docx(py_doc)
        doc.replace_tracked("Original", "Modified")

        # Get bytes (no filesystem needed, skip validation for in-memory docs)
        doc_bytes = doc.save_to_bytes(validate=False)

        # Verify bytes are valid
        assert len(doc_bytes) > 0
        assert doc_bytes[:4] == b"PK\x03\x04"

        # Can reload from bytes
        reloaded = Document(doc_bytes)
        assert "Modified" in reloaded.get_text()

    def test_multiple_edit_rounds(self, tmp_path: Path) -> None:
        """Multiple rounds of editing between libraries."""
        # Round 1: Create with python-docx
        py_doc = PythonDocxDocument()
        py_doc.add_paragraph("Version 1")

        # Round 2: Edit with python_docx_redline (skip validation for in-memory docs)
        doc = from_python_docx(py_doc)
        doc.insert_tracked(" -> Version 2", after="Version 1")
        bytes1 = doc.save_to_bytes(validate=False)

        # Round 3: Add more with python-docx
        # python-docx doesn't support loading from bytes, so save to file first
        temp_file = tmp_path / "temp.docx"
        with open(temp_file, "wb") as f:
            f.write(bytes1)
        py_doc2 = PythonDocxDocument(str(temp_file))
        py_doc2.add_paragraph("Added in round 3")
        py_doc2.save(str(temp_file))

        # Round 4: More tracked changes
        doc2 = Document(temp_file)
        doc2.insert_tracked(" [final]", after="round 3")
        doc2.save(temp_file)

        # Verify final state
        final = Document(temp_file)
        text = final.get_text()
        assert "Version 1" in text
        assert "Version 2" in text
        assert "round 3" in text
        assert "[final]" in text


class TestEdgeCases:
    """Test edge cases in python-docx integration."""

    def test_empty_document(self) -> None:
        """Can convert empty python-docx document."""
        py_doc = PythonDocxDocument()

        doc = from_python_docx(py_doc)

        # Should succeed, even if no content
        assert doc is not None

    def test_document_with_tables(self) -> None:
        """Document with tables converts correctly."""
        py_doc = PythonDocxDocument()
        py_doc.add_paragraph("Before table")
        table = py_doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        py_doc.add_paragraph("After table")

        doc = from_python_docx(py_doc)

        assert "Before table" in doc.get_text()
        assert "After table" in doc.get_text()

    def test_document_with_formatting(self) -> None:
        """Document with text formatting converts correctly."""
        py_doc = PythonDocxDocument()
        para = py_doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True
        para.add_run(" and ")
        run2 = para.add_run("italic text")
        run2.italic = True

        doc = from_python_docx(py_doc)

        text = doc.get_text()
        assert "Bold text" in text
        assert "italic text" in text
