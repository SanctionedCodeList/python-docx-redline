#!/usr/bin/env python3
"""
Reproduction script for nested <w:del> validation error.

This script demonstrates the bug where deleting an already-deleted paragraph
creates invalid nested <w:del> elements.

Usage:
    python reproduce_nested_del_bug.py

Requirements:
    - python-docx-redline
    - A test document with multiple paragraphs

The script will:
1. Create a test document (or use existing)
2. Delete a paragraph with tracked changes
3. Save successfully
4. Reload and try to delete the same paragraph again
5. Fail with ValidationError due to nested <w:del>
"""

import os
import sys
import tempfile

from lxml import etree

from python_docx_redline import Document


def check_deletion_state(doc, ref):
    """Check if a paragraph is already marked as deleted."""
    root = doc.xml_root
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    body = root.find(".//w:body", nsmap)
    paragraphs = body.findall("w:p", nsmap)

    idx = int(ref.split(":")[1])
    if idx >= len(paragraphs):
        return {"error": f"Paragraph {ref} not found"}

    p = paragraphs[idx]

    # Check for deleted paragraph mark
    ppr = p.find("w:pPr", nsmap)
    has_del_mark = False
    if ppr is not None:
        rpr = ppr.find("w:rPr", nsmap)
        if rpr is not None:
            has_del_mark = rpr.find("w:del", nsmap) is not None

    # Count deletion elements
    dels = p.findall(".//w:del", nsmap)

    return {
        "ref": ref,
        "has_deleted_paragraph_mark": has_del_mark,
        "deletion_element_count": len(dels),
        "text_content": doc.get_text_at_ref(ref),
        "is_already_deleted": has_del_mark or len(dels) > 0,
    }


def get_paragraph_xml(doc, ref):
    """Get XML string for a paragraph."""
    root = doc.xml_root
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    body = root.find(".//w:body", nsmap)
    paragraphs = body.findall("w:p", nsmap)

    idx = int(ref.split(":")[1])
    if idx >= len(paragraphs):
        return None

    return etree.tostring(paragraphs[idx], pretty_print=True).decode()


def main():
    # Check for existing test document or use provided path
    if len(sys.argv) > 1:
        test_doc_path = sys.argv[1]
        if not os.path.exists(test_doc_path):
            print(f"Error: {test_doc_path} not found")
            return 1
    else:
        # Create a simple test document
        from docx import Document as PythonDocxDocument

        test_doc_path = tempfile.mktemp(suffix=".docx")
        pydoc = PythonDocxDocument()
        for i in range(10):
            pydoc.add_paragraph(
                f"This is paragraph {i}. It contains some test content for deletion testing."
            )
        pydoc.save(test_doc_path)
        print(f"Created test document: {test_doc_path}")

    target_ref = "p:5"

    print("\n" + "=" * 60)
    print("STEP 1: Load document and check initial state")
    print("=" * 60)

    doc = Document(test_doc_path)
    initial_state = check_deletion_state(doc, target_ref)
    print(f"Initial state of {target_ref}:")
    for k, v in initial_state.items():
        print(f"  {k}: {v}")

    print("\n" + "=" * 60)
    print("STEP 2: Delete paragraph with tracked changes")
    print("=" * 60)

    doc.delete_ref(target_ref, track=True, author="Test Author")
    print(f"Deleted {target_ref}")

    after_delete = check_deletion_state(doc, target_ref)
    print("State after deletion:")
    for k, v in after_delete.items():
        print(f"  {k}: {v}")

    print("\n" + "=" * 60)
    print("STEP 3: Save document (should succeed)")
    print("=" * 60)

    try:
        doc.save(test_doc_path)
        print("✅ Save succeeded")
    except Exception as e:
        print(f"❌ Save failed: {e}")
        return 1

    print("\n" + "=" * 60)
    print("STEP 4: Reload and check state from disk")
    print("=" * 60)

    doc = Document(test_doc_path)
    reloaded_state = check_deletion_state(doc, target_ref)
    print("State after reload:")
    for k, v in reloaded_state.items():
        print(f"  {k}: {v}")

    print("\nXML structure (first 1000 chars):")
    xml = get_paragraph_xml(doc, target_ref)
    print(xml[:1000] if xml else "No XML found")

    print("\n" + "=" * 60)
    print("STEP 5: Try to delete the SAME paragraph again")
    print("=" * 60)

    print(f"Attempting to delete already-deleted {target_ref}...")
    doc.delete_ref(target_ref, track=True, author="Test Author")

    double_delete_state = check_deletion_state(doc, target_ref)
    print("State after double-deletion (in memory):")
    for k, v in double_delete_state.items():
        print(f"  {k}: {v}")

    print("\nXML structure after double-deletion (first 1500 chars):")
    xml = get_paragraph_xml(doc, target_ref)
    print(xml[:1500] if xml else "No XML found")

    print("\n" + "=" * 60)
    print("STEP 6: Try to save (should fail with ValidationError)")
    print("=" * 60)

    try:
        doc.save(test_doc_path)
        print("❌ UNEXPECTED: Save succeeded (bug not reproduced)")
        return 1
    except Exception as e:
        print(f"✅ EXPECTED: Save failed with {type(e).__name__}")
        print(f"Error message: {str(e)[:300]}")

    print("\n" + "=" * 60)
    print("REPRODUCTION COMPLETE")
    print("=" * 60)
    print("""
Root Cause:
- delete_ref() wraps content in <w:del> elements
- When called on already-deleted content, it creates nested <w:del>
- Nested <w:del> is invalid OOXML

Fix Options:
1. Check if content is already deleted before wrapping in <w:del>
2. Skip deletion if text_content is empty and deletion markers exist
3. Return early with warning/error for already-deleted content
""")

    # Cleanup
    if len(sys.argv) <= 1:
        os.remove(test_doc_path)
        print(f"Cleaned up test document: {test_doc_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
