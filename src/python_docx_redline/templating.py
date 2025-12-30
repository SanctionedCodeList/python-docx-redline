"""Document templating with markdown support.

This module provides a minimal API for creating Word documents with markdown
content. Designed for use in agent-driven document generation workflows.

Example:
    >>> from python_docx_redline.templating import DocxBuilder
    >>>
    >>> doc = DocxBuilder()
    >>> doc.heading("My Report")
    >>> doc.markdown("This is **bold** and *italic*.")
    >>> doc.table(["Name", "Value"], [["A", "1"], ["B", "2"]])
    >>> doc.save("report.docx")

Template pattern:
    >>> @dataclass
    ... class ReportData:
    ...     title: str
    ...     summary: str  # markdown
    ...
    >>> def render_report(data: ReportData, path: str) -> Path:
    ...     doc = DocxBuilder()
    ...     doc.heading(data.title)
    ...     doc.markdown(data.summary)
    ...     return doc.save(path)

Dependencies:
    This module requires optional dependencies:
        pip install python-docx-redline[templating]

    Or install directly:
        pip install markdown-it-py htmldocx
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt

if TYPE_CHECKING:
    from docx.table import _Cell


class DocxBuilder:
    """Minimal API for creating Word documents with markdown support.

    Provides a fluent interface for building documents with headings,
    paragraphs, markdown content, and tables. All methods return self
    for optional chaining.

    Args:
        landscape: Use landscape orientation (default: portrait)
        font: Font name (default: "Times New Roman")
        font_size: Font size in points (default: 11)
        margins: Page margins in inches (default: 1.0)

    Example:
        >>> doc = DocxBuilder()
        >>> doc.heading("Title")
        >>> doc.markdown("Content with **bold**")
        >>> doc.save("output.docx")

        >>> # With options
        >>> doc = DocxBuilder(landscape=True, font="Arial", font_size=12)
    """

    def __init__(
        self,
        landscape: bool = False,
        font: str = "Times New Roman",
        font_size: int = 11,
        margins: float = 1.0,
    ) -> None:
        self._doc = Document()
        self._font = font
        self._font_size = Pt(font_size)
        self._margins = Inches(margins)
        self._landscape = landscape
        self._configure()
        self._md_renderer: Any = None
        self._html_converter: Any = None

    def _configure(self) -> None:
        """Configure document layout and default styles."""
        # Set default font
        style = self._doc.styles["Normal"]
        style.font.name = self._font
        style.font.size = self._font_size

        # Configure page layout
        section = self._doc.sections[0]
        section.left_margin = self._margins
        section.right_margin = self._margins
        section.top_margin = self._margins
        section.bottom_margin = self._margins

        if self._landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11)
            section.page_height = Inches(8.5)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

    def _ensure_markdown(self) -> None:
        """Lazy-load markdown dependencies."""
        if self._md_renderer is None:
            try:
                from htmldocx import HtmlToDocx
                from markdown_it import MarkdownIt
            except ImportError as e:
                raise ImportError(
                    "Markdown support requires additional dependencies. "
                    "Install with: pip install python-docx-redline[templating] "
                    "or: pip install markdown-it-py htmldocx"
                ) from e

            self._md_renderer = MarkdownIt()
            self._html_converter = HtmlToDocx()

    def heading(self, text: str, level: int = 1) -> DocxBuilder:
        """Add a heading to the document.

        Args:
            text: Heading text
            level: Heading level (0=Title, 1=Heading1, 2=Heading2, etc.)

        Returns:
            Self for chaining
        """
        self._doc.add_heading(text, level=level)
        return self

    def paragraph(self, text: str) -> DocxBuilder:
        """Add a plain text paragraph.

        Args:
            text: Paragraph text (no formatting)

        Returns:
            Self for chaining
        """
        self._doc.add_paragraph(text)
        return self

    def markdown(self, text: str) -> DocxBuilder:
        """Add markdown content to the document.

        Supports standard markdown syntax:
        - # Headings (all levels)
        - **bold** and *italic*
        - - bullet lists and 1. numbered lists
        - > blockquotes
        - `code` and code blocks
        - | tables |

        Args:
            text: Markdown-formatted text

        Returns:
            Self for chaining

        Example:
            >>> doc.markdown('''
            ... ## Section Title
            ...
            ... This is **bold** and *italic*.
            ...
            ... - Item 1
            ... - Item 2
            ... ''')
        """
        self._ensure_markdown()
        html = self._md_renderer.render(text)
        self._html_converter.add_html_to_document(html, self._doc)
        return self

    def markdown_cell(self, cell: _Cell, text: str) -> DocxBuilder:
        """Add markdown content to a table cell.

        Args:
            cell: Table cell to add content to
            text: Markdown-formatted text

        Returns:
            Self for chaining
        """
        self._ensure_markdown()

        # Clear existing paragraphs
        for para in list(cell.paragraphs):
            para._element.getparent().remove(para._element)

        html = self._md_renderer.render(text)
        self._html_converter.add_html_to_document(html, cell)
        return self

    def table(
        self,
        headers: list[str],
        rows: list[list[str]],
        style: str = "Table Grid",
    ) -> DocxBuilder:
        """Add a table with headers and data rows.

        Args:
            headers: Column header texts
            rows: List of rows, each row is a list of cell values
            style: Table style name (default: "Table Grid")

        Returns:
            Self for chaining

        Example:
            >>> doc.table(
            ...     ["Name", "Value"],
            ...     [["Alpha", "1"], ["Beta", "2"]]
            ... )
        """
        table = self._doc.add_table(rows=1, cols=len(headers))
        table.style = style

        # Header row with bold text
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for run in header_cells[i].paragraphs[0].runs:
                run.bold = True

        # Data rows
        for row_data in rows:
            row = table.add_row()
            for i, value in enumerate(row_data):
                if i < len(row.cells):
                    row.cells[i].text = str(value)

        return self

    def table_from(
        self,
        items: list[dict[str, Any] | Any],
        columns: list[str],
        headers: list[str] | None = None,
        style: str = "Table Grid",
    ) -> DocxBuilder:
        """Create a table from a list of dicts or objects.

        Extracts the specified columns from each item and creates
        a table. Works with dicts, dataclasses, Pydantic models,
        or any object with the specified attributes.

        Args:
            items: List of dicts or objects
            columns: Attribute/key names to extract
            headers: Display headers (default: auto-generated from column names)
            style: Table style name (default: "Table Grid")

        Returns:
            Self for chaining

        Example:
            >>> @dataclass
            ... class LineItem:
            ...     description: str
            ...     quantity: int
            ...     unit_price: float
            ...
            >>> items = [
            ...     LineItem("Widget", 10, 5.00),
            ...     LineItem("Gadget", 3, 15.00),
            ... ]
            >>> doc.table_from(items, ["description", "quantity", "unit_price"])
            # Creates table with headers: Description | Quantity | Unit Price
        """
        # Generate headers from column names if not provided
        if headers is None:
            headers = [col.replace("_", " ").title() for col in columns]

        # Extract row data
        rows: list[list[str]] = []
        for item in items:
            row: list[str] = []
            for col in columns:
                if isinstance(item, dict):
                    value = item.get(col, "")
                else:
                    value = getattr(item, col, "")
                row.append(str(value))
            rows.append(row)

        return self.table(headers, rows, style=style)

    def page_break(self) -> DocxBuilder:
        """Insert a page break.

        Returns:
            Self for chaining
        """
        self._doc.add_page_break()
        return self

    def save(self, path: str | Path) -> Path:
        """Save the document to a file.

        Args:
            path: Output file path

        Returns:
            Path to the saved file
        """
        path = Path(path)
        self._doc.save(str(path))
        return path

    @property
    def document(self) -> Document:
        """Access the underlying python-docx Document.

        Use this for advanced customization not covered by
        the builder API.

        Returns:
            The python-docx Document object
        """
        return self._doc
