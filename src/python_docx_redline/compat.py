"""
Compatibility helpers for integrating with other libraries.

This module provides functions to seamlessly integrate python_docx_redline with
other document processing libraries like python-docx.
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    pass  # python-docx types imported at runtime

from .author import AuthorIdentity
from .document import Document


def from_python_docx(
    python_docx_doc: Any,
    author: str | AuthorIdentity = "Claude",
) -> Document:
    """Create a python_docx_redline Document from a python-docx Document.

    This enables workflows where you create or modify documents with python-docx
    and then add tracked changes with python_docx_redline, without needing to save
    to disk in between.

    Args:
        python_docx_doc: A python-docx Document object
        author: Author name for tracked changes (default: "Claude").
               Can also be an AuthorIdentity for full MS365 integration.

    Returns:
        A python_docx_redline Document ready for tracked change operations

    Raises:
        ImportError: If python-docx is not installed (with helpful message)
        TypeError: If the input is not a python-docx Document

    Example:
        >>> from docx import Document as PythonDocxDocument
        >>> from python_docx_redline import Document
        >>> from python_docx_redline.compat import from_python_docx
        >>>
        >>> # Create with python-docx
        >>> py_doc = PythonDocxDocument()
        >>> py_doc.add_heading("Contract", 0)
        >>> py_doc.add_paragraph("Payment terms: 30 days")
        >>>
        >>> # Convert to python_docx_redline for tracked edits
        >>> doc = from_python_docx(py_doc)
        >>> doc.replace_tracked("30 days", "45 days")
        >>> doc.save("contract_redlined.docx")

    Example with in-memory round-trip:
        >>> from docx import Document as PythonDocxDocument
        >>> from python_docx_redline.compat import from_python_docx
        >>>
        >>> # Create document
        >>> py_doc = PythonDocxDocument()
        >>> py_doc.add_paragraph("Original text")
        >>>
        >>> # Add tracked changes
        >>> doc = from_python_docx(py_doc)
        >>> doc.replace_tracked("Original", "Modified")
        >>>
        >>> # Get as bytes for storage/transmission
        >>> doc_bytes = doc.save_to_bytes()
    """
    # Runtime check for python-docx
    try:
        from docx.document import Document as PythonDocxDocType
    except ImportError as e:
        raise ImportError(
            "python-docx is required for from_python_docx(). "
            "Install it with: pip install python-docx"
        ) from e

    # Validate input type - check against the actual Document class
    if not isinstance(python_docx_doc, PythonDocxDocType):
        raise TypeError(
            f"Expected python-docx Document, got {type(python_docx_doc).__name__}. "
            "Pass a Document object created with: from docx import Document"
        )

    # Save python-docx document to BytesIO
    buffer = io.BytesIO()
    python_docx_doc.save(buffer)
    buffer.seek(0)

    # Create python_docx_redline Document from the in-memory bytes
    return Document(buffer, author=author)


def to_python_docx(doc: Document, validate: bool = True) -> Any:
    """Convert a python_docx_redline Document back to a python-docx Document.

    This is useful when you need python-docx's document creation features
    after making tracked changes with python_docx_redline.

    Args:
        doc: A python_docx_redline Document
        validate: Whether to run OOXML validation (default: True).
                 Set to False for in-memory documents without an original file.

    Returns:
        A python-docx Document object

    Raises:
        ImportError: If python-docx is not installed

    Example:
        >>> from python_docx_redline import Document
        >>> from python_docx_redline.compat import to_python_docx
        >>>
        >>> doc = Document("contract.docx")
        >>> doc.replace_tracked("old", "new")
        >>>
        >>> # Convert back to python-docx for additional operations
        >>> py_doc = to_python_docx(doc)
        >>> py_doc.add_paragraph("Added with python-docx")
        >>> py_doc.save("final.docx")
    """
    try:
        from docx import Document as PythonDocxDoc
    except ImportError as e:
        raise ImportError(
            "python-docx is required for to_python_docx(). Install it with: pip install python-docx"
        ) from e

    # Save python_docx_redline document to bytes
    doc_bytes = doc.save_to_bytes(validate=validate)

    # Load into python-docx
    buffer = io.BytesIO(doc_bytes)
    return PythonDocxDoc(buffer)
